import docx
import re


class Paper():
    """ This class handles the following:
        - process docx
        - analyse text
        - calculate score (min 1 - max 5) for each metric """

    # Headers for output CSV.
    # The filename is the id, the rest are scoring-metrics.
    FILENAME_HEADER = "filename"
    CHAPTERS_HEADER = "headings"
    INTRODUCTION_HEADER = "introduction"
    CONTENTS_HEADER = "contents"
    EPILOGUE_HEADER = "epilogue"
    SOURCES_LIST_HEADER = "sources-list"
    WORD_COUNT_HEADER = "word-count"

    # Calls constructor for docx-object
    def __init__(self, filePath):
        # Creating word reader object.
        try:
            self.doc = docx.Document(filePath)
        except IOError:
            raise
        self.chapterHeadings = []
        # get full text of docx in a string
        self.fullText = ""
        for paragraph in self.doc.paragraphs:
            self.fullText += paragraph.text

    # Returns list of CSV-headers
    @classmethod
    def getCSVHeaders(cls):
        return [cls.FILENAME_HEADER,
                cls.CHAPTERS_HEADER,
                cls.INTRODUCTION_HEADER,
                cls.CONTENTS_HEADER,
                cls.EPILOGUE_HEADER,
                cls.SOURCES_LIST_HEADER,
                cls.WORD_COUNT_HEADER]    

    # Returns dictionary with all computed scores.     
    def getScores(self):
        scores = {}

        self.chapterHeadings = self.getChapterHeadings()

        # See if we can identify the presence of special headings.
        # Its metric is scored, then the heading is removed from the list.
        scores[self.INTRODUCTION_HEADER] = self.scoreAndPopMatch("inleiding|voorwoord")
        scores[self.CONTENTS_HEADER] = self.scoreAndPopMatch("inhoud(s)?(opgave)?|hoofdstukken")
        scores[self.EPILOGUE_HEADER] = self.scoreAndPopMatch("nawoord|conclusie")
        scores[self.SOURCES_LIST_HEADER] = self.scoreAndPopMatch("bronnen(lijst)?")

        # The headings we have left are the chapter headings, score the count.
        scores[self.CHAPTERS_HEADER] = self.getChapterHeadingsScore()
        
        # Score the word count, the assignment said 1000-2000 words.
        scores[self.WORD_COUNT_HEADER] = self.getWordCountScore()

        print(self.chapterHeadings)
        return scores

    # Returns list of all chapter-headings found in docx.
    # Chapter headings are identified by word-style.
    def getChapterHeadings(self):
        chapterHeadings = []

        for paragraph in self.doc.paragraphs:
            if re.match("heading", paragraph.style.name, re.IGNORECASE):
                chapterHeadings.append(paragraph.text)
        return chapterHeadings

    # If we get a match in the header score=5, else score=1
    # The match is removed from the headings-list
    def scoreAndPopMatch(self, regexMatchString):
        for heading in self.chapterHeadings:
            if re.match(regexMatchString, heading, re.IGNORECASE):
                self.chapterHeadings.remove(heading)
                return 5
        return 1
    
    # Returns score (1-5) for the number of chapters.     
    def getChapterHeadingsScore(self):
        chapterCount = len(self.chapterHeadings)
        score = 1

        # The sweet spot is 4-8 chapters
        if 4 <= chapterCount <= 8:
            score = 5
        # 3 chapters is low, but acceptable
        elif chapterCount == 3:
            score = 3
        # between 9 and 12 seems excessive, but acceptable
        elif 9 <= chapterCount <= 12:
            score = 3
        # The rest is either too few or too many
        else: score = 1

        return score

    # Returns score based on deviation from prescribed word count
    def getWordCountScore(self):
        wordCount = len(self.fullText.strip().split(' '))

        score = 1
        # The sweet spot is 1000-2000 words
        if 1000 <= wordCount <= 2000:
            score = 5
        # Less than minimal is unacceptable
        elif wordCount < 1000:
            score = 1
        # More than max is okay, up to a point
        elif 2000 < wordCount <=3000:
            score = 3
        # If they make me read too much, it's not appreciated.
        else: 
            score = 1
        return score